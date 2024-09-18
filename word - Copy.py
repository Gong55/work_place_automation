from docx import Document
from docx.text.paragraph import Paragraph
from docx.text.run import Run
import os
import pandas as pd
import re

excel_file = r'E:\Desktop\python\Word\contract_input.xlsx'
template_file = r'E:\Desktop\python\Word\SIC Draft Contract.docx'
destination_folder = r'E:\Desktop\python\Word\output'

df = pd.read_excel(excel_file)
df = df.fillna(' ')
df = df.T


def format_room(room):
    if re.fullmatch(r"\d+/\d+", room):
        room = re.sub(r"(\d+)/(\d+)", r"\1-\2", room)
        room = f"({room})"
    else:
        room = re.sub(r"(\d+)/(\d+)", r"\1-\2", room)
    return room


def find_and_replace_text(variable, old_text, new_text):
    if old_text in variable.text:
        variable.text = variable .text.replace(old_text, new_text)


for index, row in df.iterrows():
    doc = Document(template_file)

    for section in doc.sections:
        footer = section.footer
        for footer_text in footer.paragraphs:
            find_and_replace_text(
                footer_text, "{owner_name}", str(row['owner_name']))

    for paragraph in doc.paragraphs:
        find_and_replace_text(
            paragraph, "{owner_name}", str(row['owner_name']))
        find_and_replace_text(
            paragraph, "{room_number}", str(row['room_number']))

    formatted_room = format_room(str(row['room_number']))
    new_file_name = os.path.join(
        destination_folder, f"Lease Agreement {row['project_name']} {formatted_room}.docx")
    doc.save(new_file_name)


print(f"'output_{formatted_room}.docx' created successfully")
