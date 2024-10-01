from docx import Document
from docx.text.paragraph import Paragraph
from docx.text.run import Run
import os
import pandas as pd
import re
from datetime import datetime
import numpy as np
from dateutil.relativedelta import relativedelta
from pythainlp.util import bahttext
from num2words import num2words


base_path = r'E:\Desktop\python\Word'
excel_file = os.path.join(base_path, 'data_input.xlsx')
lease_agreement = os.path.join(base_path, 'SIC Draft Contract.docx')
destination_folder = os.path.join(base_path, 'output')
owner_contract = os.path.join(base_path, 'owner_contract.docx')

df = pd.read_excel(excel_file)
df_flipped = df.set_index('Attributes').transpose()

thai_month_names = {
    1: 'มกราคม',
    2: 'กุมภาพันธ์',
    3: 'มีนาคม',
    4: 'เมษายน',
    5: 'พฤษภาคม',
    6: 'มิถุนายน',
    7: 'กรกฎาคม',
    8: 'สิงหาคม',
    9: 'กันยายน',
    10: 'ตุลาคม',
    11: 'พฤศจิกายน',
    12: 'ธันวาคม'
}


def number_to_text_en(df, column):
    df[column +
        '_text_en'] = df[column].apply(lambda x: num2words(x) + ' Baht only',)


def calculate_two_months_deposit(df, column):
    df[column + '_times_two'] = df[column].apply(lambda x: x * 2)


def number_to_text_th(df, column):
    df[column + '_text_th'] = df[column].apply(bahttext)


def lease_period(start_date, end_date, df):
    duration_list = []
    for start, end in zip(df[start_date], df[end_date]):
        start = start + pd.DateOffset(days=1)
        duration = relativedelta(end, start)

        components = []

        if duration.years:
            components.append(f"{duration.years} Year" +
                              ("s" if duration.years > 1 else ""))
        if duration.months:
            components.append(f"{duration.months} Month" +
                              ("s" if duration.months > 1 else ""))
        if duration.days:
            components.append(f"{duration.days} Day" +
                              ("s" if duration.days > 1 else ""))

        formatted_duration = ", ".join(components)

        duration_list.append(formatted_duration)

    return duration_list


def late_payment_grace_period(start_date, df):

    answer = []

    for start in df[start_date]:
        output = start + relativedelta(days=4)
        answer.append(output.days, output.months, output.years)
    return answer


def convert_date_format(column, df):
    df[column] = df[column].replace(' ', np.nan)
    df[column] = pd.to_datetime(
        df[column], errors='coerce')
    if df[column].notna().all():
        df[f'{column}_day'] = df[column].dt.day
        df[f'{column}_month'] = df[column].dt.month
        df[f'{column}_year'] = df[column].dt.year
        df[f'{column}_month_en'] = df[column].dt.month_name()
        df[f'{column}_month_th'] = df[f'{column}_month'].map(
            thai_month_names)
        df[column + '_year_th'] = df[column + '_year'].apply(lambda x: x + 543)
        df[column + '_en'] = df[column + '_day'].astype(
            str) + ' ' + df[column + '_month_en'] + ' ' + df[column + '_year'].astype(str)
        df[column + '_th'] = df[column + '_day'].astype(
            str) + ' ' + df[column + '_month_th'] + ' ' + df[column + '_year'].astype(str)
    else:
        pass


def format_room(room):
    if re.fullmatch(r"\d+/\d+", room):
        room = re.sub(r"(\d+)/(\d+)", r"\1-\2", room)
        room = f"({room})"
    else:
        room = re.sub(r"(\d+)/(\d+)", r"\1-\2", room)
    return room


convert_date_format("rent_start_date", df_flipped)
convert_date_format("rent_end_date", df_flipped)
convert_date_format("owner_passport_expire_date", df_flipped)
convert_date_format("owner_passport_expire_date_2", df_flipped)
convert_date_format("tenant_passport_expire_date", df_flipped)
convert_date_format("tenant_passport_expire_date_2", df_flipped)


def replace_text_with_format(paragraph, old_text, new_text):
    new_text = str(new_text)
    for run in paragraph.runs:
        if old_text in run.text:
            if old_text.isupper():
                new_text = new_text.upper()
            elif old_text.islower():
                new_text = new_text.lower()
            elif old_text.istitle():
                new_text = new_text.title()

            run.text = run.text.replace(old_text, new_text)


def replace_text_in_tables(table, old_text, new_text):
    for table_row in table.rows:
        for cell_row in table_row.cells:
            if cell_row.paragraphs:
                for cell_paragraph in cell_row.paragraphs:
                    replace_text_with_format(
                        cell_paragraph, old_text, new_text if pd.notna(new_text) else '')


def replace_text_if_df_exist(paragraph, old_text, new_text):
    replace_text_with_format(
        paragraph, old_text, new_text if pd.notna(new_text) else '')


for index, row in df_flipped.iterrows():
    doc = Document(owner_contract)

    for paragraph in doc.paragraphs:
        replace_text_with_format(
            paragraph, 'PROJECTNAMEHOLDER', row['project_name'])
        replace_text_with_format(
            paragraph, 'UNITNUMBERHOLDER', row['room_number'])
        replace_text_with_format(
            paragraph, 'STYAHD', row['rent_start_date_year'])
        replace_text_with_format(
            paragraph, 'STMDHD', row['rent_start_date_month'])
        replace_text_with_format(
            paragraph, 'STDYHD', row['rent_start_date_day'])
        replace_text_with_format(
            paragraph, 'ENYAHD', row['rent_end_date_year'])
        replace_text_with_format(
            paragraph, 'ENMDHD', row['rent_end_date_month'])
        replace_text_with_format(
            paragraph, 'ENDYHD', row['rent_end_date_day'])
    for table in doc.tables:
        replace_text_in_tables(table, 'NAME1HOLDER', row['owner_name'])
        replace_text_in_tables(table, 'NAME2HOLDER',
                               f"/ {row['owner_name_2']}")
        replace_text_in_tables(
            table, 'NAME2NOSLASHHOLDER', row['owner_name_2'])
        replace_text_in_tables(table, 'PSPT1HO', row['owner_passport'])
        replace_text_in_tables(
            table, 'PSPT2HO', f"/ {row['owner_passport_2']}")

    file_path = os.path.join(destination_folder, 'Test.docx')
    iteration = 1
    while os.path.exists(file_path):
        file_path = os.path.join(
            destination_folder, f"Test_{str(iteration)}.docx")
        iteration += 1

    doc.save(file_path)
    print(f"File has been saved to {file_path}")

# paragraph_placeholder_dictionary = {
#     'Startdateplahor': 'start_date_en',
#     'Startdatethplahor': 'start_date_th',
#     'OWNERNAMEPLAHOR': 'owner_name',
#     'OWNERPASSPORTPLAHOR': 'owner_passport',
#     'Ownernationalityplahor': 'owner_nationality',
#     'Ownerpassportexpiredateplahor': 'owner_passport_expire_date_en',
#     'Projectnameplahor': 'project_name',
#     'Roomnumberplahor': 'room_number',
#     'Buildingnumberplahor': 'building_no',
#     'Projectaddressplahor': 'project_address',
#     'Ownernationalitythplahor': 'owner_nationality_th',
#     'Ownerpassportexpiredatethplahor': 'owner_passport_expire_date_th',
#     'Projectaddressthplahor': 'project_address_th',
#     'TENANTNAMEPLAHOR': 'tenamt_name',
#     'Tenantpassportplahor': 'tenant_passport',
#     'Tenantnationalityplahor': 'tenant_nationality',
#     'Tenantnationalitythplahor': 'tenant_nationality_th',
#     'Tenantpassportexpiredateplahor': 'owner_passport_expire_date_en',
#     'Tenantpassportexpiredatethplahor': 'owner_passport_expire_date_th',
#     'Floorplahor': 'room_floor_number',
#     'Areaplahor': 'room_area',
#     'Rentnoplahor': 'rent_price',
#     'Rentenplahor': 'rent_price_text_en',
#     'Rentthplahor': 'rent_price_text_th',
#     'Enddateplahor': 'end_date_th',
#     'Depositplahor': 'rent_price_times_two',
#     # 'Leaseperioud': '' 1 Year, 6 months
#     'Delosittextplahor': 'rent_price_times_two_text_en',
#     'Delosittextthplahor': 'rent_price_times_two_text_th',
#     'Startdayplahor': 'start_date_day',
#     # day before fine
#     'OWNERBANKNAMEPLAHOR': 'owner_bank',
#     'Ownerbankbranch': 'owner_bank_branch',
#     'Ownerbankaccountno': 'owner_bank_account_no',
#     'OWNERBANKACCOUNTNAMEPLAHOR': 'owner_bank_account_name',
#     'Watermeternoplahor': 'water_meter',
#     'electricmeternoplahor': 'electric_meter'


# }

# footer_placeholder_dictionary = {
#     'OWNERNAMEPLAHOR': 'owner_name',
#     'TENANTNAMEPLAHOR': 'tenamt_name',
#     'WITNESSNAMEPLAHOR': 'witness_name',


# }


# def process_document(doc, data_row, placeholders, dictionary):
#     for paragraph in doc.paragraphs:
#         for placeholder, column_name in dictionary.item():
#             e


# for index2, row2 in df_flipped.iterrows():
#     doc2 = Document(lease_agreement)

#     for paragraph in doc2.paragraphs:
#         for placeholder, column_name in paragraph_placeholder_dictionary.items():
#             value = str(row2[column_name]) if pd.notna(
#                 row2[column_name]) else ''
#             replace_text_with_format(paragraph, placeholder, value)

#     for section in doc.sections:
#         footer = section.footer
#         for footer_text in footer.paragraphs:

# for index, row in df_flipped.iterrows():
#     doc = Document(template_file)

#     for section in doc.sections:
#         footer = section.footer
#         for footer_text in footer.paragraphs:
#             replace_text_with_format(
#                 footer_text, "owner_name", str(row['owner_name']))
#             replace_text_with_format(
#                 footer_text, "owner_name_2", str(row['owner_name_2']))
#             replace_text_with_format(
#                 footer_text, "tenant_name", str(row['tenant_name']))
#             replace_text_with_format(
#                 footer_text, "tenant_name_2", str(row['tenant_name_2']))
#             replace_text_with_format(
#                 footer_text, "witness_name", str(row['witness_name']))

#     for paragraph in doc.paragraphs:

#         replace_text_with_format(
#             paragraph, "project_name", str(paragraph['project_name']))
#         replace_text_with_format(
#             paragraph, "-room_number-", str(row['room_number']))
#         replace_text_with_format(
#             paragraph, "floor_number", str(row['floor_number']))
#         replace_text_with_format(paragraph, "area", str(row['area']))
#         replace_text_with_format(
#             paragraph, "building_no", str(row['building_no']))
#         replace_text_with_format(paragraph, "address", str(row['address']))
#         # address TH
#         # start_date
#         # end_date
#         # start_date TH
#         # end_date TH
#         replace_text_with_format(
#             paragraph, "rent_price", str(row['rent_price']))
#         # rent_price EN
#         # rent_price TH
#         # rent_price x2
#         # rent_price x2 EN
#         # rent_price x2 TH

#         replace_text_with_format(paragraph, "late_payment_grace_period", str(
#             row['late_payment_grace_period']))
#         #
#         replace_text_with_format(
#             paragraph, "-owner_name-", str(row['owner_name']))
#         replace_text_with_format(
#             paragraph, "owner_passport", str(['owner_passport']))
#         replace_text_with_format(
#             paragraph, "owner_nationality", str(['owner_nationality']))
#         replace_text_with_format(paragraph, "owner_bank", str(['owner_bank']))
#         replace_text_with_format(
#             paragraph, "owner_bank_branch", str('owner_bank_branch'))
#         replace_text_with_format(
#             paragraph, "owner_bank_account_no", str(['owner_bank_account_no']))
#         replace_text_with_format(
#             paragraph, "owner_bank_account_name", str(['owner_bank_account_name']))

#         # owner_passport_expire_date
#         # owner_passport_expire_date TH

#     formatted_room = format_room(str(row['room_number']))
#     new_file_name = os.path.join(
#         destination_folder, f"Lease Agreement {row['project_name']} {formatted_room}.docx")
#     doc.save(new_file_name)

# for index, row in df_flipped.iterrows():
#     doc2 = Document(owner_contract)
#     for table in doc2.tables:
#         replace_text_in_tables(table, "-NAME1-", str(row['owner_name']))
#         replace_text_in_tables(table, "-NANE2-", str(row['owner_name_2']))
#         replace_text_in_tables(table, "-PASPRT1-", str(row['owner_passport']))
#         replace_text_in_tables(
#             table, "-PASPRT1-", str(row['pwner_passport_2']))

#     for paragraph2 in doc2.paragraphs:
#         replace_text_with_format(
#             paragraph2, "-UNITNAME-", str(row['project_name']))
#         replace_text_with_format(
#             paragraph2, "-UNITNUMBER-", str(row['room_number']))
#         # replace_text_with_format(paragraph2, "-SY-", )
#         # replace_text_with_format(paragraph2, "-SM-", )
#         # replace_text_with_format(paragraph2, "-SD-", )
#         # replace_text_with_format(paragraph2, "-ENY-", )
#         # replace_text_with_format(paragraph2, "-ENM-", )
#         # replace_text_with_format(paragraph2, "-EMD-", )
#         # Have to convert excel date into DD MM YY

#     formatted_room = format_room(str(row['room_number']))
#     new_file_name = os.path.join(
#         destination_folder, f"代租管合約 {row['project_name']} {formatted_room}.docx")
#     doc2.save(new_file_name)
#     print(f"'output_{formatted_room}.docx' created successfully")
